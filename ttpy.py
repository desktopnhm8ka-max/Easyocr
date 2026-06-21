import os
import csv
import tempfile
import fitz  # PyMuPDF
import easyocr
from docx import Document
from tkinter import Tk
from tkinter.filedialog import askopenfilename, asksaveasfilename

# ---------------- SETTINGS ----------------
ROW_THRESHOLD = 15
OCR_DPI = 200

# ---------------- TABLE RECONSTRUCTION ----------------
def extract_table_from_ocr(result):
    words = []

    for bbox, text, confidence in result:
        x = bbox[0][0]
        y = bbox[0][1]

        words.append({
            "x": x,
            "y": y,
            "text": text.strip()
        })

    # Sort top-to-bottom then left-to-right
    words.sort(key=lambda w: (w["y"], w["x"]))

    rows = []

    for word in words:
        matched = False

        for row in rows:
            if abs(row["y"] - word["y"]) <= ROW_THRESHOLD:
                row["items"].append(word)
                matched = True
                break

        if not matched:
            rows.append({
                "y": word["y"],
                "items": [word]
            })

    table_rows = []

    for row in rows:
        row["items"].sort(key=lambda w: w["x"])

        row_texts = [
            item["text"]
            for item in row["items"]
            if item["text"]
        ]

        if row_texts:
            table_rows.append(row_texts)

    return table_rows


# ---------------- MAIN ----------------
print("===== PDF TABLE OCR =====")

Tk().withdraw()

pdf_path = askopenfilename(
    title="Select PDF",
    filetypes=[("PDF Files", "*.pdf")]
)

if not pdf_path:
    print("No PDF selected.")
    exit()

print("Loading OCR Engine...")
reader = easyocr.Reader(['en'], gpu=False)

all_rows = []

try:

    with fitz.open(pdf_path) as pdf:

        total_pages = len(pdf)

        print(f"Total Pages: {total_pages}")

        for page_num in range(total_pages):

            try:

                percent = ((page_num + 1) / total_pages) * 100

                print(
                    f"Processing Page "
                    f"{page_num + 1}/{total_pages} "
                    f"({percent:.1f}%)"
                )

                page = pdf.load_page(page_num)

                # ------------------------------------------------
                # FIRST TRY NORMAL PDF TEXT EXTRACTION
                # ------------------------------------------------
                page_text = page.get_text().strip()

                if page_text:

                    print("  Text PDF detected")

                    for line in page_text.splitlines():

                        line = line.strip()

                        if line:
                            all_rows.append([line])

                    continue

                # ------------------------------------------------
                # OCR FALLBACK
                # ------------------------------------------------
                print("  OCR mode")

                pix = page.get_pixmap(
                    dpi=OCR_DPI,
                    alpha=False
                )

                with tempfile.NamedTemporaryFile(
                    suffix=".png",
                    delete=False
                ) as temp_img:

                    temp_image_path = temp_img.name

                pix.save(temp_image_path)

                result = reader.readtext(temp_image_path)

                os.remove(temp_image_path)

                page_rows = extract_table_from_ocr(result)

                all_rows.extend(page_rows)

            except Exception as page_error:

                print(
                    f"Page {page_num + 1} skipped:\n"
                    f"{page_error}"
                )

                continue

except Exception as pdf_error:

    print("Failed to open PDF:")
    print(pdf_error)
    exit()

# ---------------- NO DATA CHECK ----------------
if not all_rows:
    print("No text could be extracted.")
    exit()

print("\nExtraction Completed Successfully")
print(f"Rows Extracted: {len(all_rows)}")

# ---------------- SAVE CSV ----------------
csv_path = asksaveasfilename(
    title="Save CSV File",
    defaultextension=".csv",
    filetypes=[("CSV Files", "*.csv")]
)

if csv_path:

    max_cols = max(len(row) for row in all_rows)

    with open(
        csv_path,
        "w",
        newline="",
        encoding="utf-8"
    ) as file:

        writer = csv.writer(file)

        for row in all_rows:

            padded = row + [""] * (max_cols - len(row))

            writer.writerow(padded)

    print("CSV Saved Successfully")

# ---------------- SAVE TXT ----------------
txt_path = asksaveasfilename(
    title="Save TXT File",
    defaultextension=".txt",
    filetypes=[("Text Files", "*.txt")]
)

if txt_path:

    with open(
        txt_path,
        "w",
        encoding="utf-8"
    ) as file:

        for row in all_rows:
            file.write(" | ".join(row) + "\n")

    print("TXT Saved Successfully")

# ---------------- SAVE DOCX ----------------
docx_path = asksaveasfilename(
    title="Save DOCX File",
    defaultextension=".docx",
    filetypes=[("Word Documents", "*.docx")]
)

if docx_path:

    doc = Document()

    doc.add_heading("PDF OCR Table Output", 0)

    max_cols = max(len(row) for row in all_rows)

    table = doc.add_table(
        rows=len(all_rows),
        cols=max_cols
    )

    table.style = "Table Grid"

    for r, row in enumerate(all_rows):

        for c, value in enumerate(row):

            table.cell(r, c).text = str(value)

    doc.save(docx_path)

    print("DOCX Saved Successfully")

print("\nProject Completed Successfully")   