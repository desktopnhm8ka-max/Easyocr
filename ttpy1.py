import os
import csv
import tempfile
import fitz  # PyMuPDF
import easyocr
from docx import Document
from docx.shared import Pt
from tkinter import Tk
from tkinter.filedialog import askopenfilename

# ---------------- SETTINGS ----------------
ROW_THRESHOLD = 15
OCR_DPI = 200

# Supported image formats (matches report: "PDF and Image Text Extraction")
IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp")


# ================================================================
# MODULE 1 — TABLE RECONSTRUCTION
# ================================================================
def extract_table_from_ocr(result):
    """
    Reconstructs table-like structures from EasyOCR output.
    Words with similar Y-axis positions are grouped into rows,
    then sorted by X-axis position within each row.
    """
    words = []

    for bbox, text, confidence in result:
        x = bbox[0][0]
        y = bbox[0][1]
        words.append({"x": x, "y": y, "text": text.strip()})

    words.sort(key=lambda w: (w["y"], w["x"]))

    rows = []

    for word in words:
        matched = False
        for row in rows:
            if abs(row["y"] - word["y"]) <= ROW_THRESHOLD:
                row["items"].append(word)
                matched = True
                break
        if not matched:
            rows.append({"y": word["y"], "items": [word]})

    table_rows = []

    for row in rows:
        row["items"].sort(key=lambda w: w["x"])
        row_texts = [item["text"] for item in row["items"] if item["text"]]
        if row_texts:
            table_rows.append(row_texts)

    return table_rows


# ================================================================
# MODULE 2 — OCR PROCESSING
# ================================================================
def run_ocr_on_image(reader, image_path):
    """
    Passes an image file to EasyOCR and returns reconstructed table rows.
    """
    result = reader.readtext(image_path)
    return extract_table_from_ocr(result)


def extract_from_pdf(pdf_path, reader):
    """
    Processes a PDF file page by page.
    - Pages with embedded text: direct extraction via PyMuPDF (fast, accurate).
    - Pages with no text (scanned): converted to image and processed via OCR.
    """
    all_rows = []

    with fitz.open(pdf_path) as pdf:
        total_pages = len(pdf)
        print(f"Total Pages: {total_pages}")

        for page_num in range(total_pages):
            try:
                print(f"Processing Page {page_num + 1}/{total_pages}")
                page = pdf.load_page(page_num)
                text = page.get_text().strip()

                if text:
                    # Direct text extraction — no OCR needed
                    for line in text.splitlines():
                        line = line.strip()
                        if line:
                            all_rows.append([line])
                    continue

                # No readable text found — switch to OCR mode
                print(f"  → No text found, switching to OCR mode")
                pix = page.get_pixmap(dpi=OCR_DPI, alpha=False)

                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    img_path = tmp.name

                pix.save(img_path)

                page_rows = run_ocr_on_image(reader, img_path)
                all_rows.extend(page_rows)

                # Clean up temporary image file
                os.remove(img_path)

            except Exception as e:
                print(f"  ✗ Page {page_num + 1} failed: {e}")
                continue

    return all_rows


def extract_from_image(image_path, reader):
    """
    Processes a standalone image file directly through OCR.
    """
    print(f"Processing image: {os.path.basename(image_path)}")
    return run_ocr_on_image(reader, image_path)


# ================================================================
# MODULE 3 — CSV OUTPUT
# ================================================================
def save_csv(all_rows, output_path):
    """
    Saves extracted data in CSV format — spreadsheet-compatible.
    Useful for data analysis and record keeping.
    """
    max_cols = max(len(r) for r in all_rows)

    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for row in all_rows:
            writer.writerow(row + [""] * (max_cols - len(row)))

    print(f"CSV Saved:  {output_path}")


# ================================================================
# MODULE 4 — TXT OUTPUT
# ================================================================
def save_txt(all_rows, output_path):
    """
    Saves extracted data as plain text — easy to read and edit.
    """
    with open(output_path, "w", encoding="utf-8") as f:
        for row in all_rows:
            f.write(" | ".join(row) + "\n")

    print(f"TXT Saved:  {output_path}")


# ================================================================
# MODULE 5 — DOCX OUTPUT
# ================================================================
def save_docx(all_rows, output_path):
    """
    Saves extracted data as a Microsoft Word document with
    reconstructed table structure — suitable for reporting and
    documentation purposes.
    """
    doc = Document()
    doc.add_heading("OCR Output", 0)

    if not all_rows:
        doc.add_paragraph("No content extracted.")
        doc.save(output_path)
        return

    # Determine number of columns from the widest row
    num_cols = max(len(row) for row in all_rows)

    # Build a proper table in the Word document
    table = doc.add_table(rows=len(all_rows), cols=num_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(all_rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            # Consistent font across all cells
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.save(output_path)
    print(f"DOCX Saved: {output_path}")


# ================================================================
# MAIN
# ================================================================
print("===== PDF AND IMAGE TEXT EXTRACTION =====")

# --- File Selection ---
Tk().withdraw()

pdf_path = askopenfilename(
    title="Select a PDF or Image File",
    filetypes=[
        ("Supported Files", "*.pdf *.png *.jpg *.jpeg *.bmp *.tiff *.tif *.webp"),
        ("PDF Files", "*.pdf"),
        ("Image Files", "*.png *.jpg *.jpeg *.bmp *.tiff *.tif *.webp"),
    ]
)

if not pdf_path:
    print("No file selected.")
    exit()

file_ext = os.path.splitext(pdf_path)[1].lower()

# --- Load OCR Engine ---
print("Loading OCR Engine...")
reader = easyocr.Reader(['en'], gpu=False)

# --- Extract Content ---
all_rows = []

try:
    if file_ext == ".pdf":
        all_rows = extract_from_pdf(pdf_path, reader)
    elif file_ext in IMAGE_EXTENSIONS:
        all_rows = extract_from_image(pdf_path, reader)
    else:
        print(f"Unsupported file type: {file_ext}")
        exit()
except Exception as e:
    print(f"Extraction failed: {e}")
    exit()

if not all_rows:
    print("No text extracted.")
    exit()

print(f"\nExtraction Completed Successfully")
print(f"Rows Extracted: {len(all_rows)}")

# --- Auto Save ---
output_folder = os.path.dirname(pdf_path)
base_name = os.path.splitext(os.path.basename(pdf_path))[0]

csv_path  = os.path.join(output_folder, base_name + "_OCR.csv")
txt_path  = os.path.join(output_folder, base_name + "_OCR.txt")
docx_path = os.path.join(output_folder, base_name + "_OCR.docx")

print()
save_csv(all_rows, csv_path)
save_txt(all_rows, txt_path)
save_docx(all_rows, docx_path)

print("\nProject Completed Successfully")
